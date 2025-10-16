import os
import time
import json 
from google import genai
from google.genai.errors import APIError
from docx import Document 
from typing import List, Dict, Any 
from datetime import datetime
import shutil
# --- API ANAHTARLARI VE İSTEMCİ YÖNETİMİ ---
API_KEYS = [
    # Bu kısmı doldurmalısınız!
    "AIzaSyAdcdT89H3bCOvDYAtMxhAv29QW3UWUsjo",
    "AIzaSyCrsTVZ7xNh6flE6gQDtHS9n1Ky919VqLc",
]

class GeminiClientManager:
    """Çoklu API anahtarlarını yönetir ve Resource Exhausted hatalarında anahtar değiştirir."""
    def __init__(self, api_keys):
        self.api_keys = api_keys
        self.current_key_index = 0
        self.client = self._initialize_client()

    def _initialize_client(self):
        """Mevcut dizindeki anahtarla bir Gemini istemcisi başlatır."""
        if not self.api_keys:
            raise ValueError("API anahtar listesi boş olamaz.")
        
        current_key = self.api_keys[self.current_key_index]
        print(f"\n[CLIENT] Yeni anahtar yükleniyor (Index: {self.current_key_index + 1}/{len(self.api_keys)}).")
        try:
            return genai.Client(api_key=current_key)
        except Exception as e:
            raise Exception(f"Gemini istemcisi başlatılırken hata: {e}")

    def switch_client(self):
        """Sıradaki API anahtarına geçer ve istemciyi yeniden başlatır."""
        self.current_key_index += 1
        if self.current_key_index >= len(self.api_keys):
            print("\n[KRİTİK HATA] Tüm API anahtarları tükenmiştir veya limitlerine ulaşmıştır.")
            raise StopIteration("Tüm API anahtarları tükendi.")
            
        self.client = self._initialize_client()
        return self.client

    def get_client(self):
        """Mevcut istemciyi döndürür."""
        return self.client

# --- BAŞLANGIÇ KONTROLLERİ ---
print(">>> BETİK BAŞLADI! Kütüphaneler başarıyla yüklendi. <<<")

try:
    client_manager = GeminiClientManager(API_KEYS)
except StopIteration as e:
    print(f"Betik durduruldu: {e}")
    exit(1)
except Exception as e:
    print(f"Başlangıç hatası: {e}")
    exit(1)


# Yapılandırma
MODEL_NAME = 'gemini-2.5-flash'  
SOURCE_LANG = 'English'
TARGET_LANG = 'Turkish'
MAX_CHUNK_TOKENS = 5000 
DELAY_BETWEEN_CALLS = 1 

# --- DOCX METİN YÖNETİMİ ---

def get_all_text_holders(doc):
    """Belgedeki tüm metin tutucularını (Paragraf ve Hücre içi Paragraflar) döndürür."""
    holders = []
    # 1. Ana Paragraflar
    holders.extend(doc.paragraphs)
    # 2. Tablo Hücrelerindeki Paragraflar (Tablolar dahil)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                holders.extend(cell.paragraphs)
    return holders

def extract_text_chunks_from_docx(docx_path: str) -> tuple[List[str], Dict[str, Any], Document]:
    """DOCX'ten metinleri çıkarır."""
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise Exception(f"DOCX Dosya Açma Hatası: '{docx_path}'. Belge bozuk olabilir. Hata: {e}")

    chunk_data = {} 
    
    for holder in get_all_text_holders(doc):
        original_text = holder.text.strip() 
        if not original_text:
            continue
            
        if original_text not in chunk_data:
            chunk_data[original_text] = {
                'translated': None, 
                'holders': []
            }
        chunk_data[original_text]['holders'].append(holder)
            
    text_chunks = list(chunk_data.keys())
                    
    return text_chunks, chunk_data, doc

def chunk_text_by_tokens(text_chunks: List[str], max_tokens: int, client_mgr: GeminiClientManager) -> List[List[str]]:
    """Metin listesini token limitine uygun Batch'lere böler."""
    
    batches = []
    current_batch = []
    current_token_count = 0
    
    for text in text_chunks:
        try:
            token_count = client_mgr.get_client().models.count_tokens(
                model=MODEL_NAME,
                contents=[text]
            ).total_tokens
        except Exception:
            token_count = len(text) // 4
            
        chunk_cost = token_count + 10
        
        if current_token_count + chunk_cost > max_tokens:
            if current_batch:
                batches.append(current_batch)
            current_batch = [text]
            current_token_count = chunk_cost
        else:
            current_batch.append(text)
            current_token_count += chunk_cost
            
    if current_batch:
        batches.append(current_batch)
        
    return batches

def translate_batch(text_batch: List[str], client_mgr: GeminiClientManager) -> List[str]:
    """Toplu metin parçalarını çevirir ve hata kurtarma yapar."""
    
    full_text = json.dumps(text_batch, ensure_ascii=False)
    
    prompt = (
        f"Aşağıdaki {len(text_batch)} adet metin dizisini ('{SOURCE_LANG}' dilinde) Türkçe'ye çevir. "
        "Yanıtı, aynı sırada çevrilmiş metinlerden oluşan bir JSON listesi olarak döndür. "
        "Çeviri sırasında orijinal metinlerdeki tüm özel karakterleri ve etiketleri (örneğin: {VARIABLE} veya [TAG]) kesinlikle koru. "
        "ÇEVİRİ TALİMATLARI: Kod bloklarını, Assembly komutlarını ve Komut Satırı çıktılarını **ASLA ÇEVİRME** ve orijinal haliyle KORU. "
        f"Çevrilecek Metin Dizisi: {full_text}"
    )
    
    max_retries_api = 5 
    while True: # Başarana kadar tekrar dene
        for api_attempt in range(max_retries_api):
            try:
                current_client = client_mgr.get_client()
                response = current_client.models.generate_content(
                    model=MODEL_NAME,
                    contents=prompt,
                    config=genai.types.GenerateContentConfig(response_mime_type="application/json")
                )
                
                raw_response_text = response.text.strip()
                if raw_response_text.startswith("```json"):
                    raw_response_text = raw_response_text[7:]
                if raw_response_text.endswith("```"):
                    raw_response_text = raw_response_text[:-3]

                translated_texts = json.loads(raw_response_text.strip())
                
                if len(translated_texts) != len(text_batch):
                    raise ValueError("Liste boyutları eşleşmiyor.")
                
                return translated_texts
                
            except (APIError, json.JSONDecodeError, ValueError) as e:
                is_resource_exhausted = 'Resource Exhausted' in str(e) or '429' in str(e) or '400' in str(e)
                
                if is_resource_exhausted and client_mgr.current_key_index < len(client_mgr.api_keys) - 1:
                    print(f"\n[API KOTA HATASI] Anahtar {client_mgr.current_key_index + 1} kotası doldu.")
                    client_mgr.switch_client()
                    time.sleep(5) 
                    continue 
                
                # 🚨 DÜZELTME: 'attempt' yerine 'api_attempt' kullanıldı.
                wait_time = 2 ** api_attempt
                if api_attempt < max_retries_api - 1:
                    print(f"\n[HATA] Deneme {api_attempt + 1}: Hata: {type(e).__name__}. {wait_time} saniye bekleniyor...")
                    time.sleep(wait_time) 
                    continue
                else:
                    raise Exception(f"Kalıcı Hata ({type(e).__name__}) - Batch Atlanamadı.")
            except StopIteration:
                raise
            except Exception as e:
                raise Exception(f"Beklenmedik Hata: {e}")

def update_docx_with_translations(doc: Document, translation_map: Dict[str, str], docx_path: str):
    """Orijinal DOCX dosyasını çevrilmiş metinlerle biçimlendirme korumalı olarak günceller."""
    
    output_path = docx_path 
    
    # 🚨 BİÇİMLENDİRME KORUMALI GÜNCELLEME
    for holder in get_all_text_holders(doc):
        original_text = holder.text.strip()
        
        if original_text in translation_map:
            translated_text = translation_map[original_text]
            
            if not holder.runs:
                holder.text = translated_text
                continue
                
            first_run = holder.runs[0]
            
            for run in reversed(holder.runs):
                run.clear()

            first_run.text = translated_text
            
    # Orijinal dosyanın üzerine yazar
    doc.save(output_path)
    return output_path

def automate_docx_translation(input_filepath, output_filepath):
    """DOCX çeviri sürecini yönetir."""
    
    # 🚨 GÜVENLİK ADIMI: Orijinal dosyanın üzerine yazmadan önce yedeklemesini al
    backup_filepath = f"{input_filepath}.bak.{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    if not os.path.exists(backup_filepath):
        shutil.copyfile(input_filepath, backup_filepath)
        print(f"[YEDEKLEME] Orijinal dosya '{backup_filepath}' olarak yedeklendi.")
    
    # 1. Metinleri çıkar ve Geçici Kayıt Dosyasını Yönet
    print(f"'{input_filepath}' dosyasından metin çıkarılıyor...")
    
    try:
        all_text_chunks, chunk_data_map, doc = extract_text_chunks_from_docx(input_filepath)
    except Exception as e:
        raise Exception(f"Metin Çıkarım Hatası: {e}")
    
    base_file_name, ext = os.path.splitext(input_filepath)
    temp_translation_file = f"{base_file_name}_temp_translations.json"
    
    # Çevirileri depolayacak harita
    translated_map = {}
    
    if os.path.exists(temp_translation_file):
        try:
            with open(temp_translation_file, 'r', encoding='utf-8') as f:
                translated_map = json.load(f)
            print(f"[AKILLI YÜKLEME] {len(translated_map)} adet önceden çevrilmiş parça bulundu.")
        except json.JSONDecodeError:
            print("[UYARI] Geçici kayıt dosyası bozuk, sıfırdan başlanıyor.")
    
    # 2. Çevrilmemiş parçaları filtrele
    untranslated_chunks = [chunk for chunk in all_text_chunks if chunk not in translated_map]
    
    if not untranslated_chunks:
        print("\n*** Çevrilecek yeni parça bulunamadı. Orijinal dosya güncelleniyor... ***")
        # Final güncellemesi (zaten çevrilmiş olanları dosyanın üzerine yazar)
        # 🚨 DÜZELTME: update_docx_with_translations, translation_map bekler.
        final_output_path = update_docx_with_translations(doc, translated_map, input_filepath)
        print(f"İşlem tamamlandı. Orijinal dosya '{final_output_path}' olarak güncellendi.")
        return

    # 3. Batch'lere böl
    batches = chunk_text_by_tokens(untranslated_chunks, MAX_CHUNK_TOKENS, client_manager)
    total_batches = len(batches)
    
    print(f"Toplam {total_batches} Batch çevrilecek. Kalan tahmini süre: ~{(total_batches * (DELAY_BETWEEN_CALLS + 5)) // 60} dakika.")
    
    # 4. Çeviri Döngüsü
    for batch_num, batch in enumerate(batches):
        print(f"\n[İŞLENİYOR] Batch {batch_num + 1}/{total_batches} ({len(batch)} parça) çevriliyor...")
        
        try:
            # Çeviri fonksiyonu
            translated_texts = translate_batch(batch, client_manager) 
            
            # 5. Haritayı Güncelle ve Geçici Kayıt
            for original, translated in zip(batch, translated_texts):
                translated_map[original] = translated
            
            # Her başarılı Batch'te geçici dosyayı kaydet
            with open(temp_translation_file, 'w', encoding='utf-8') as f:
                json.dump(translated_map, f, ensure_ascii=False, indent=2)
                
            print(f"[BAŞARILI KAYIT] Batch {batch_num + 1} kaydedildi.")
            time.sleep(DELAY_BETWEEN_CALLS)
            
        except StopIteration:
            raise
        except Exception as e:
            # Kritik hata oluşursa (API hatası dahil) dışarı fırlatılır.
            print(f"\n[KRİTİK HATA] Batch {batch_num + 1} çevirisi başarısız oldu. Hata: {e}")
            raise # Ana döngüye fırlat
        
    # 6. Final DOCX Güncellemesi
    # Orijinal dosyanın üzerine yazar
    final_output_path = update_docx_with_translations(doc, translated_map, input_filepath)

    print(f"\nİşlem tamamlandı. Çevrilmiş DOCX '{final_output_path}' olarak kaydedildi.")

# --- KULLANIM ---
if __name__ == "__main__":
    INPUT_DOCX = 'girdi_belgesi.docx' 
    OUTPUT_DOCX = 'girdi_belgesi_TR.docx' 
    
    crash_count = 0
    
    while True:
        try:
            if crash_count > 0:
                print("\n\n#####################################################")
                print(f"!!! YENİDEN BAŞLATILIYOR (Tekrar Deneme Sayısı: {crash_count}) !!!")
                print("#####################################################")
                time.sleep(5) 
            
            # Ana çeviri fonksiyonunu çağır
            automate_docx_translation(INPUT_DOCX, OUTPUT_DOCX)
            
            print("\n\n*** ÇEVİRİ İŞLEMİ BAŞARIYLA TAMAMLANDI! ***")
            break
            
        except StopIteration:
            print("\n\n*** TÜM API ANAHTARLARI TÜKENDİ. BETİK SONLANDIRILDI. ***")
            break
            
        except Exception as e:
            crash_count += 1
            print(f"\n[ANA KONTROL] Program kritik bir hata ile durdu: {type(e).__name__} ({e}). Yeniden başlatılıyor...")
            time.sleep(5)